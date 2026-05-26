"""Build DOCX exports from the canonical markdown docs.

Run from repo root:

    python scripts/build_docs.py

Outputs land in ``docs/exports/``. Re-run after editing any source markdown.

Deck PDFs are produced from the HTML decks in ``docs/exports/decks/`` via
headless Chrome/Edge — the @page rules in deck.css already encode
1280×720, zero margins, with print backgrounds.
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "exports"
OUT.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(0x1A, 0x2A, 0x4A)
GREY = RGBColor(0x55, 0x5C, 0x6E)
BLACK = RGBColor(0x10, 0x10, 0x14)


# ---------- markdown → docx (intentionally minimal: headings, paragraphs, lists) ----------


def md_to_docx(md_path: Path, docx_path: Path, title: str) -> None:
    """Convert a markdown file to a styled Word document.

    Handles: H1/H2/H3, paragraphs, unordered lists, ordered lists, simple tables,
    code/horizontal rules, italics-only blockquotes. Not a general-purpose converter
    — calibrated for the Kairos doc style.
    """
    doc = Document()

    # default style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    md = md_path.read_text(encoding="utf-8")
    lines = md.splitlines()

    in_table: list[list[str]] = []

    def flush_table() -> None:
        if not in_table:
            return
        rows = in_table[:]
        in_table.clear()
        if len(rows) < 2:
            return
        # rows[1] is the |---|---| separator; skip it
        header = [c.strip() for c in rows[0]]
        body = [[c.strip() for c in r] for r in rows[2:]]
        table = doc.add_table(rows=1 + len(body), cols=len(header))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(header):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        for ri, row in enumerate(body, start=1):
            for ci, val in enumerate(row):
                if ci < len(table.rows[ri].cells):
                    table.rows[ri].cells[ci].text = val

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # table?
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table.append([c for c in stripped.strip("|").split("|")])
            i += 1
            continue
        else:
            flush_table()

        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped):
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # headings
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[2:])
            r.font.size = Pt(20)
            r.font.bold = True
            r.font.color.rgb = NAVY
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:])
            r.font.size = Pt(16)
            r.font.bold = True
            r.font.color.rgb = NAVY
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:])
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = BLACK
            i += 1
            continue

        # blockquote (treat as italic grey)
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = GREY
            i += 1
            continue

        # unordered list
        if re.match(r"^[-*+] ", stripped):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, text)
            i += 1
            continue

        # ordered list
        if re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, text)
            i += 1
            continue

        # paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    flush_table()
    doc.save(docx_path)


_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def _add_inline_runs(paragraph, text: str) -> None:
    """Render bold/italic/code/link inline markers as Word runs."""
    parts = _INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.font.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            r.font.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = NAVY
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            label = part[1 : part.index("](")]
            r = paragraph.add_run(label)
            r.font.color.rgb = NAVY
            r.font.underline = True
        else:
            paragraph.add_run(part)


# ---------- driver ----------


DOC_TARGETS = [
    ("docs/MASTER.md", "Kairos — Master documentation index", "00_master_index.docx"),
    ("docs/ONBOARDING.md", "Kairos — Contractor onboarding", "01_contractor_onboarding.docx"),
    ("docs/business/strategy.md", "Kairos — Commercialization strategy", "02_strategy.docx"),
    ("docs/business/state.md", "Kairos — Technical state", "03_technical_state.docx"),
    ("docs/external/one-pager.md", "Kairos — Sales one-pager", "04_sales_one_pager.docx"),
    ("docs/external/security.md", "Kairos — Security & data handling", "05_security_overview.docx"),
    ("docs/investor/pitch-outline.md", "Kairos — Pitch outline", "06_pitch_outline.docx"),
    ("docs/ROADMAP.md", "Kairos — Commercialization roadmap", "08_roadmap.docx"),
]

DECK_TARGETS = [
    ("docs/exports/decks/pitch.html", "07_pitch_deck.pdf"),
    ("docs/exports/decks/roadmap.html", "09_roadmap_deck.pdf"),
]


# ---------- HTML deck → PDF (via headless Chrome/Edge) ----------


def _find_chromium() -> str | None:
    """Return a path to a Chromium-based browser that supports --print-to-pdf."""
    env = os.environ.get("KAIROS_CHROME")
    if env and Path(env).exists():
        return env

    candidates = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    for name in ("google-chrome", "chromium", "chromium-browser", "msedge"):
        found = shutil.which(name)
        if found:
            return found
    return None


def html_to_pdf(html_path: Path, pdf_path: Path, browser: str) -> None:
    """Render an HTML deck to PDF using headless Chromium."""
    file_url = "file:///" + str(html_path.resolve()).replace("\\", "/")
    cmd = [
        browser,
        "--headless=new",
        "--disable-gpu",
        "--no-margins",
        "--print-to-pdf-no-header",
        f"--print-to-pdf={pdf_path}",
        file_url,
    ]
    subprocess.run(cmd, check=True, capture_output=True)


def main() -> None:
    print(f"Output: {OUT}")
    for src_rel, title, out_name in DOC_TARGETS:
        src = ROOT / src_rel
        if not src.exists():
            print(f"  skip (missing) {src_rel}")
            continue
        dst = OUT / out_name
        md_to_docx(src, dst, title)
        print(f"  wrote {out_name}")

    browser = _find_chromium()
    if browser is None:
        print("  skip decks: no Chrome/Edge found (set KAIROS_CHROME to override)")
        return
    for src_rel, out_name in DECK_TARGETS:
        src = ROOT / src_rel
        if not src.exists():
            print(f"  skip (missing) {src_rel}")
            continue
        dst = OUT / out_name
        html_to_pdf(src, dst, browser)
        print(f"  wrote {out_name}")


if __name__ == "__main__":
    main()
